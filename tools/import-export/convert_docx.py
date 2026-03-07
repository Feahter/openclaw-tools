#!/usr/bin/env python3
from docx import Document
import sys

doc = Document('/Users/fuzhuo/.openclaw/workspace/assets/离线版文字游戏设计与开发最佳实践指南.docx')

output = []
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    
    style = para.style.name if para.style else ''
    if 'Heading' in style:
        level = style.replace('Heading ', '')
        try:
            level = int(level)
            output.append(f"{'#' * level} {text}")
        except:
            output.append(f"## {text}")
    else:
        output.append(text)

for table in doc.tables:
    output.append('')
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        output.append('| ' + ' | '.join(cells) + ' |')

with open('/Users/fuzhuo/.openclaw/workspace/assets/离线版文字游戏设计与开发最佳实践指南.md', 'w', encoding='utf-8') as f:
    f.write('\n\n'.join(output))

print('转换完成')
